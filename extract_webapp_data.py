#!/usr/bin/env python3
"""
Script pour extraire toutes les données de la WebApp Neskao Trade Desk
et générer un rapport Word complet (.docx)
"""

import json
import re
from datetime import datetime
from pathlib import Path
import sys

# Vérifier si python-docx est installé
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("⚠️  Module python-docx non trouvé. Installation...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

class NeskaoDataExtractor:
    def __init__(self, webapp_dir):
        self.webapp_dir = Path(webapp_dir)
        self.sections_dir = self.webapp_dir / "src" / "components" / "sections"
        self.data = {}
        
    def extract_data_from_file(self, file_path):
        """Extrait les données structurées d'un fichier React/TypeScript"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Extraire les données des arrays et objets
            data_patterns = [
                r'const\s+(\w+)\s*=\s*\[(.*?)\];',
                r'const\s+(\w+)\s*=\s*\{(.*?)\};',
                r'(\w+Data)\s*=\s*\[(.*?)\];',
                r'(\w+Data)\s*=\s*\{(.*?)\};'
            ]
            
            extracted = {}
            for pattern in data_patterns:
                matches = re.finditer(pattern, content, re.DOTALL)
                for match in matches:
                    var_name = match.group(1)
                    var_content = match.group(2)
                    if len(var_content) > 50:  # Seulement les gros objets de données
                        extracted[var_name] = var_content[:500] + "..." if len(var_content) > 500 else var_content
            
            return extracted
        except Exception as e:
            print(f"Erreur lecture {file_path}: {e}")
            return {}
    
    def extract_all_data(self):
        """Extrait toutes les données de tous les composants"""
        print("🔍 Extraction des données de la WebApp...")
        
        sections = [
            "Dashboard.tsx", "Contexte.tsx", "Reglementation.tsx", 
            "Produits.tsx", "Financement.tsx", "SGA.tsx", 
            "Rentabilite.tsx", "ImpactSocial.tsx", "AnalyseDecisionnelle.tsx",
            "Risques.tsx", "NextSteps.tsx"
        ]
        
        for section in sections:
            section_file = self.sections_dir / section
            if section_file.exists():
                section_name = section.replace('.tsx', '')
                self.data[section_name] = self.extract_data_from_file(section_file)
                print(f"  ✅ {section_name}: {len(self.data[section_name])} objets extraits")
        
        # Extraire aussi les données des fichiers de données
        data_files = [
            "rentabilite-updateddatas.js",
            "financement-updateddatas.txt", 
            "impactsocial-updateddatas.txt"
        ]
        
        for data_file in data_files:
            data_path = self.webapp_dir / "reports" / "datas tsx" / data_file
            if data_path.exists():
                self.data[data_file] = self.extract_data_from_file(data_path)
        
        return self.data

class NeskaoReportGenerator:
    def __init__(self, data):
        self.data = data
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Configure les styles du document"""
        # Style titre principal
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.name = 'Calibri'
        
        # Style sous-titre
        subtitle_style = self.doc.styles.add_style('CustomSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.bold = True
        subtitle_style.font.name = 'Calibri'
        
        # Style texte normal
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
    
    def add_header(self):
        """Ajoute l'en-tête du rapport"""
        # Page de garde
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("NESKAO TRADE DESK\nÉTUDE DE FAISABILITÉ")
        run.font.size = Pt(20)
        run.font.bold = True
        
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Analyse Comparative de Localisation\nBureau de Trading International")
        run.font.size = Pt(14)
        
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_p.add_run(f"Rapport généré le {datetime.now().strftime('%d/%m/%Y')}")
        run.font.size = Pt(12)
        run.font.italic = True
        
        self.doc.add_page_break()
        
        # Table des matières
        toc = self.doc.add_paragraph("TABLE DES MATIÈRES", style='CustomTitle')
        toc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        sections_toc = [
            "1. CONTEXTE ET ENVIRONNEMENT",
            "2. MÉTHODOLOGIE D'ANALYSE", 
            "3. ANALYSE RÉGLEMENTAIRE",
            "4. MIX PRODUITS ET STRATÉGIE",
            "5. STRUCTURE FINANCIÈRE",
            "6. ANALYSE SG&A",
            "7. RENTABILITÉ ET ROI",
            "8. IMPACT SOCIAL",
            "9. ANALYSE DÉCISIONNELLE",
            "10. GESTION DES RISQUES",
            "11. NEXT STEPS ET ROADMAP",
            "12. CONCLUSIONS ET RECOMMANDATIONS"
        ]
        
        for section in sections_toc:
            p = self.doc.add_paragraph(section)
            p.style = 'List Number'
        
        self.doc.add_page_break()
    
    def add_executive_summary(self):
        """Ajoute le résumé exécutif"""
        self.doc.add_heading('RÉSUMÉ EXÉCUTIF', level=1)
        
        summary_text = """
Cette étude présente une analyse exhaustive de 12 localisations potentielles pour l'établissement 
du bureau de trading international de Neskao. L'analyse multicritères évalue chaque localisation 
selon 5 dimensions clés : réglementation (15%), impact social (30%), ROI (15%), 
financement DFI (10%) et cash management (30%).

RECOMMANDATION PRINCIPALE : PARIS 🇫🇷

Avec un score pondéré de 8.09/10, Paris représente le choix optimal grâce à :
• Liens historiques privilégiés avec la Côte d'Ivoire
• Accès optimisé aux financements de développement (AFD/Proparco) 
• Impact social maximum (150K diaspora ivoirienne)
• EBITDA positif dès l'An 1 (+0.52M€)
• Convention fiscale éliminant le risque de double imposition

ALTERNATIVES CRÉDIBLES :
• Genève (8.06/10) - Hub financier international, standards suisses
• Amsterdam (7.98/10) - Port #1 cacao Europe, innovation ESG

INVESTISSEMENT REQUIS : 1.89M€ de capital initial
RENTABILITÉ : ROI 3 ans de 171.6%, payback 2.3 ans
IMPACT : 50+ emplois qualifiés, formation jeunes Ivoiriens
        """
        
        self.doc.add_paragraph(summary_text.strip())
        self.doc.add_page_break()
    
    def generate_section_dashboard(self):
        """Génère la section Dashboard"""
        self.doc.add_heading('1. CONTEXTE ET ENVIRONNEMENT', level=1)
        
        # Présentation de Neskao
        self.doc.add_heading('1.1 Neskao : Pionnier Africain de la Transformation du Cacao', level=2)
        
        context_text = """
Fondée en septembre 2013, Neskao est la première entreprise africaine à transformer 
les fèves de cacao hors normes en produits semi-finis de qualité. Cette société familiale 
fondée par Jean Pierre Roux et dirigée par Sylvie Roux a développé un modèle d'affaires 
révolutionnaire qui valorise les déchets de la filière cacao.

CHIFFRES CLÉS :
• Création : 2013 (1ère en Afrique)
• Emplois : 150+ directs, 8000 indirects  
• Capacité : 32K tonnes/an
• Certification : FSSC 22000 V.5 (2021)

CAPACITÉS ACTUELLES :
• Pâte de cacao : 12 000 tonnes/an
• Beurre de cacao : 5 000 tonnes/an  
• Tourteau de cacao : 15 000 tonnes/an
• Localisation : Zone industrielle de Vridi, Abidjan
        """
        
        self.doc.add_paragraph(context_text.strip())
        
        # Environnement international
        self.doc.add_heading('1.2 Environnement International du Cacao', level=2)
        
        market_text = """
CHIFFRES MONDIAUX :
• Production mondiale : 4.5M tonnes/an
• Production Côte d'Ivoire : 2.0M tonnes/an (45%)
• Capacité transformation : 980K tonnes installée
• Horizon 2029 : 1.9M tonnes capacité

TENDANCES MAJEURES :
📈 Volatilité accrue des prix nécessitant des outils sophistiqués (forwards/futures)
🌱 Réglementation EUDR créant des opportunités pour les acteurs conformes
💹 Financiarisation du marché via ICE Futures
🤝 Consolidation des traders, fenêtre d'opportunité pour nouveaux entrants africains
        """
        
        self.doc.add_paragraph(market_text.strip())
    
    def generate_section_methodology(self):
        """Génère la section méthodologie"""
        self.doc.add_heading('2. MÉTHODOLOGIE D\'ANALYSE', level=1)
        
        methodology_text = """
L'analyse comparative des 12 localisations s'appuie sur une approche multicritères pondérée :

CRITÈRES D'ÉVALUATION :
• Réglementation (15%) : Cadre légal, conventions fiscales, restrictions
• Impact Social (30%) : Liens avec CI, diaspora, programmes formation
• ROI (15%) : Rentabilité opérationnelle, payback, IRR
• Financement DFI (10%) : Accès institutions développement, conditions
• Cash Management (30%) : Besoins liquidité, coûts financement

LOCALISATIONS ÉVALUÉES :
1. Paris 🇫🇷 - Score : 8.09/10
2. Genève 🇨🇭 - Score : 8.06/10  
3. Amsterdam 🇳🇱 - Score : 7.98/10
4. Singapour 🇸🇬 - Score : 7.49/10
5. Hambourg 🇩🇪 - Score : 7.32/10
6. Londres 🇬🇧 - Score : 7.06/10
7. Chypre 🇨🇾 - Score : 7.14/10
8. Maroc CFC 🇲🇦 - Score : 6.91/10
9. Maurice 🇲🇺 - Score : 6.56/10
10. Tel Aviv 🇮🇱 - Score : 6.58/10
11. Dubai 🇦🇪 - Score : 6.50/10
12. Andorre 🇦🇩 - Score : 5.23/10

DONNÉES SOURCES :
• Réglementations nationales et conventions fiscales
• Données ICE Futures et marchés dérivés cacao
• Analyses AFD, Proparco, IFC sur financements développement
• Études d'impact social et diaspora
• Projections financières 3 ans (2024-2026)
        """
        
        self.doc.add_paragraph(methodology_text.strip())
    
    def generate_all_sections(self):
        """Génère toutes les sections du rapport"""
        print("📝 Génération du rapport Word...")
        
        # En-tête et résumé
        self.add_header()
        self.add_executive_summary()
        
        # Sections principales
        self.generate_section_dashboard()
        self.generate_section_methodology()
        
        # TODO: Ajouter les autres sections avec les vraies données extraites
        sections_to_add = [
            ("3. ANALYSE RÉGLEMENTAIRE", "Reglementation"),
            ("4. MIX PRODUITS ET STRATÉGIE", "Produits"), 
            ("5. STRUCTURE FINANCIÈRE", "Financement"),
            ("6. ANALYSE SG&A", "SGA"),
            ("7. RENTABILITÉ ET ROI", "Rentabilite"),
            ("8. IMPACT SOCIAL", "ImpactSocial"),
            ("9. ANALYSE DÉCISIONNELLE", "AnalyseDecisionnelle"),
            ("10. GESTION DES RISQUES", "Risques"),
            ("11. NEXT STEPS ET ROADMAP", "NextSteps")
        ]
        
        for title, section_key in sections_to_add:
            self.doc.add_heading(title, level=1)
            
            if section_key in self.data:
                section_data = self.data[section_key]
                self.doc.add_paragraph(f"[Section {section_key} - Données extraites : {len(section_data)} objets]")
                
                # Ajouter un échantillon des données
                for key, value in list(section_data.items())[:3]:
                    p = self.doc.add_paragraph()
                    p.add_run(f"{key}: ").bold = True
                    p.add_run(str(value)[:200] + "..." if len(str(value)) > 200 else str(value))
            else:
                self.doc.add_paragraph("[Données en cours d'extraction...]")
            
            self.doc.add_page_break()
        
        # Conclusions
        self.doc.add_heading('12. CONCLUSIONS ET RECOMMANDATIONS', level=1)
        
        conclusion_text = """
DÉCISION STRATÉGIQUE RECOMMANDÉE : PARIS

L'analyse exhaustive des 12 localisations confirme Paris comme choix optimal pour 
l'établissement du bureau de trading international de Neskao.

FACTEURS DÉCISIFS :
✅ Score consolidé : 8.09/10 (1er/12)
✅ Impact social maximum : 8.5/10 (liens CI, diaspora 150K)
✅ Rentabilité immédiate : EBITDA +0.52M€ dès An1
✅ Financement privilégié : AFD/Proparco (score 10/10)
✅ Cadre réglementaire optimal : Convention fiscale CI

PLAN DE DÉPLOIEMENT :
Phase 1 (Août 2024) : Fondation légale & réglementaire
Phase 2 (Sept-Nov 2024) : Structuration financière & partenariats  
Phase 3 (Oct-Déc 2024) : Déploiement opérationnel

INVESTISSEMENT & RETOUR :
• Capital initial : 1.89M€
• ROI 3 ans : 171.6%
• Payback : 2.3 ans
• Résultats nets cumulés : 5.14M€

Cette recommandation positionne Neskao comme acteur pionnier du trading cacao africain 
sur les marchés internationaux, maximisant l'impact économique et social.
        """
        
        self.doc.add_paragraph(conclusion_text.strip())
    
    def save_report(self, output_path):
        """Sauvegarde le rapport"""
        print(f"💾 Sauvegarde du rapport : {output_path}")
        self.doc.save(output_path)
        print(f"✅ Rapport généré avec succès : {output_path}")

def main():
    # Configuration
    webapp_dir = Path(__file__).parent
    timestamp = datetime.now().strftime("%d.%m.%y")
    output_path = webapp_dir / "reports" / f"NESKAO_TradeDesk_Rapport_Faisabilite_{timestamp}_v1.docx"
    
    print("🚀 Génération du rapport Neskao Trade Desk...")
    print(f"📂 Répertoire WebApp : {webapp_dir}")
    print(f"📄 Fichier de sortie : {output_path}")
    
    # Extraction des données
    extractor = NeskaoDataExtractor(webapp_dir)
    data = extractor.extract_all_data()
    
    print(f"📊 Données extraites : {sum(len(v) for v in data.values())} objets total")
    
    # Génération du rapport
    generator = NeskaoReportGenerator(data)
    generator.generate_all_sections()
    
    # Sauvegarde
    output_path.parent.mkdir(exist_ok=True)
    generator.save_report(output_path)
    
    print("🎉 Génération terminée avec succès!")
    return output_path

if __name__ == "__main__":
    main()